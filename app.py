import os
import torch
import asyncio
import numpy as np
import streamlit as st
import librosa
import librosa.display
from docx import Document
from faster_whisper import WhisperModel
from sklearn.cluster import SpectralClustering
from python_speech_features import mfcc

# ----- Environment Fixes -----
os.environ["TORCH_CPU_ONLY"] = "1"
os.environ["CUDA_VISIBLE_DEVICES"] = "-1"
torch.set_default_dtype(torch.float32)

if not hasattr(asyncio, "get_running_loop"):
    asyncio.get_event_loop().close()
    asyncio.set_event_loop(asyncio.new_event_loop())

st.title("🎤 AI Audio Summarizer with Speaker Identification")

# ----- Session State Initialization -----
if "transcript_text" not in st.session_state:
    st.session_state.transcript_text = None

# ----- Upload Audio File -----
uploaded_audio = st.file_uploader("Upload Audio File (MP3, WAV, M4A)", type=["mp3", "wav", "m4a"])
if uploaded_audio is not None:
    audio_ext = uploaded_audio.name.split('.')[-1]
    audio_path = f"temp_audio.{audio_ext}"

    # Save the uploaded file to disk
    with open(audio_path, "wb") as f:
        f.write(uploaded_audio.read())
    st.success("✅ Audio Uploaded! Click 'Transcribe Audio' to process.")

    # ----- Transcribe Audio and Detect Speakers -----
    if st.button("Transcribe Audio"):
        with st.spinner("🔍 Transcribing and Detecting Speakers..."):
            try:
                # Initialize the Faster-Whisper model (CPU only)
                model = WhisperModel("small", device="cpu", compute_type="int8")
                segments, info = model.transcribe(audio_path, beam_size=5, vad_filter=True, word_timestamps=True)
                
                segment_texts = []
                start_times = []
                for segment in segments:
                    segment_texts.append(segment.text)
                    start_times.append(segment.start)

                # ----- Speaker Diarization (Lightweight Alternative) -----
                st.info("🔍 Performing Speaker Diarization...")

                # Load audio file and extract MFCC embeddings
                y, sr = librosa.load(audio_path, sr=16000)
                mfcc_features = mfcc(y, sr, numcep=13, winlen=0.025, winstep=0.01)
                
                # Use Spectral Clustering for speaker separation
                num_speakers = min(3, len(mfcc_features) // 100)
                clustering = SpectralClustering(n_clusters=num_speakers, affinity='nearest_neighbors', assign_labels='kmeans')
                speaker_labels = clustering.fit_predict(mfcc_features[:len(segment_texts)])  # Match segment count
                
                # Assign speakers to transcript
                transcript_text = ""
                speaker_mapping = {}

                for i, (text, label, start) in enumerate(zip(segment_texts, speaker_labels, start_times)):
                    speaker_id = f"Speaker {label + 1}"
                    if speaker_id not in speaker_mapping:
                        speaker_mapping[label] = f"Speaker {len(speaker_mapping) + 1}"
                    speaker_name = speaker_mapping[label]
                    
                    transcript_text += f"{speaker_name}: {text}\n\n"

                st.session_state.transcript_text = transcript_text  # Store transcript in session state

                # Save the transcript to a Word document
                transcript_doc = Document()
                transcript_doc.add_heading("Audio Transcript", level=1)
                transcript_doc.add_paragraph(transcript_text)
                transcript_path = "audio_transcript.docx"
                transcript_doc.save(transcript_path)

                st.success("✅ Transcription Complete!")
                with open(transcript_path, "rb") as f:
                    st.download_button("📥 Download Transcript", f, "audio_transcript.docx")
            except Exception as e:
                st.error(f"❌ Error in transcription: {e}")

    # ----- Summarization Function -----
    def chunk_text(text, max_chunk_size=1000):
        chunks = []
        while len(text) > max_chunk_size:
            chunk = text[:max_chunk_size]
            last_period = chunk.rfind('.')
            if last_period == -1:
                last_period = max_chunk_size
            chunk = text[:last_period+1]
            chunks.append(chunk.strip())
            text = text[last_period+1:]
        if text.strip():
            chunks.append(text.strip())
        return chunks

    # ----- Summarize Transcript -----
    if st.session_state.transcript_text and st.button("Summarize Transcript"):
        with st.spinner("📝 Summarizing..."):
            try:
                from transformers import pipeline
                summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

                chunks = chunk_text(st.session_state.transcript_text, max_chunk_size=1000)
                summaries = []
                for chunk in chunks:
                    summary = summarizer(chunk, max_length=300, min_length=100, do_sample=False)
                    summaries.append(summary[0]['summary_text'])

                final_summary = "\n\n".join(summaries)

                # Save to Word document
                summary_doc = Document()
                summary_doc.add_heading("Meeting Summary", level=1)
                summary_doc.add_paragraph(final_summary)
                summary_path = "summary.docx"
                summary_doc.save(summary_path)

                st.success("✅ Summary Generated!")
                with open(summary_path, "rb") as f:
                    st.download_button("📥 Download Summary", f, "summary.docx")
            except Exception as e:
                st.error(f"❌ Error in summarization: {e}")

    # ----- Cleanup Temporary Files -----
    for filename in [audio_path, "audio_transcript.docx", "summary.docx"]:
        if os.path.exists(filename):
            os.remove(filename)
